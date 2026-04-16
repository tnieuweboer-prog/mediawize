# modules/workbook/builder.py (IMPROVED VERSION)
"""
Workbook DOCX builder with improved error handling and memory management.
"""
from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ============================================================
# CONSTANTS
# ============================================================

STEP_LABEL = "Stap"
MATERIAALSTAAT_LABEL = "Materiaalstaat"
STAPPEN_LABEL = "Stappen"
DEFAULT_TITLE = "Werkboekje"
TITLE_FONT_SIZE = 20
STEP_HEADING_FONT_SIZE = 14
IMAGE_WIDTH_INCHES = 2.0
COVER_WIDTH_INCHES = 6.5
IMAGES_PER_ROW = 3


# ============================================================
# HELPERS
# ============================================================

def _add_title(doc: Document, title: str) -> None:
    """
    Add title paragraph to document.
    
    Args:
        doc: Document object
        title: Title text
    """
    try:
        p = doc.add_paragraph()
        run = p.add_run(title.strip() if title else DEFAULT_TITLE)
        run.bold = True
        run.font.size = Pt(TITLE_FONT_SIZE)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception as e:
        logger.error(f"Error adding title: {e}")
        raise


def _add_meta_block(doc: Document, meta: dict[str, Any]) -> None:
    """
    Add metadata table to document.
    
    Args:
        doc: Document object
        meta: Metadata dictionary
    """
    try:
        vak = (meta.get("vak") or "").upper()
        profieldeel = meta.get("profieldeel") or ""
        docent = meta.get("docent") or ""
        duur = meta.get("duur") or ""
        
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        
        def add_row(k: str, v: str) -> None:
            r = table.add_row().cells
            r[0].text = k
            r[1].text = v
        
        if vak:
            add_row("Vak", vak)
        if profieldeel:
            add_row("Profieldeel", profieldeel)
        if docent:
            add_row("Docent", docent)
        if duur:
            add_row("Duur", duur)
        
        doc.add_paragraph()  # Whitespace
    except Exception as e:
        logger.error(f"Error adding meta block: {e}")
        raise


def _try_add_image(doc: Document, img_bytes: bytes, width_inches: float) -> bool:
    """
    Safely add image to document with proper resource cleanup.
    
    Args:
        doc: Document object
        img_bytes: Image bytes
        width_inches: Image width in inches
        
    Returns:
        True if successful, False otherwise
    """
    if not img_bytes:
        logger.warning("Empty image bytes provided")
        return False
    
    bio = None
    try:
        bio = io.BytesIO(img_bytes)
        bio.seek(0)
        doc.add_picture(bio, width=Inches(width_inches))
        return True
    except Exception as e:
        logger.error(f"Error adding image: {e}")
        return False
    finally:
        if bio:
            try:
                bio.close()
            except Exception as e:
                logger.warning(f"Error closing BytesIO: {e}")


def _add_cover(doc: Document, meta: dict[str, Any]) -> None:
    """
    Add cover image to document.
    
    Args:
        doc: Document object
        meta: Metadata dictionary
    """
    cover = meta.get("cover_bytes")
    if not cover:
        return
    
    try:
        if _try_add_image(doc, cover, width_inches=COVER_WIDTH_INCHES):
            doc.add_paragraph()
        else:
            logger.warning("Failed to add cover image, continuing without it")
            doc.add_paragraph()
    except Exception as e:
        logger.error(f"Error adding cover: {e}")
        # Cover is optional, continue without it
        doc.add_paragraph()


def _add_materiaalstaat(doc: Document, meta: dict[str, Any]) -> None:
    """
    Add materials list table to document.
    
    Args:
        doc: Document object
        meta: Metadata dictionary
    """
    if not meta.get("include_materiaalstaat"):
        return
    
    try:
        materialen = meta.get("materialen") or []
        if not materialen:
            return
        
        doc.add_paragraph().add_run(MATERIAALSTAAT_LABEL).bold = True
        
        headers = ["Nummer", "Aantal", "Benaming", "Lengte", "Breedte", "Dikte", "Materiaal"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        # Add header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(h)
            run.bold = True
        
        # Add data rows
        for item in materialen:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                row_cells[i].text = str(item.get(h, "") or "")
        
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Error adding materiaalstaat: {e}")
        raise


def _add_step(doc: Document, idx: int, step: dict[str, Any]) -> None:
    """
    Add step section to document.
    
    Args:
        doc: Document object
        idx: Step index (1-based)
        step: Step dictionary
    """
    try:
        title = (step.get("title") or "").strip()
        text_blocks = step.get("text_blocks") or []
        images = step.get("images") or []
        
        # Add step heading
        heading = doc.add_paragraph()
        heading_text = f"{STEP_LABEL} {idx}: {title}" if title else f"{STEP_LABEL} {idx}"
        run = heading.add_run(heading_text)
        run.bold = True
        run.font.size = Pt(STEP_HEADING_FONT_SIZE)
        
        # Add text blocks
        for block in text_blocks:
            b = (block or "").strip()
            if b:
                doc.add_paragraph(b)
        
        # Add images in grid layout
        if images:
            cols = IMAGES_PER_ROW
            table = doc.add_table(rows=0, cols=cols)
            table.autofit = True
            
            row = None
            for i, img_bytes in enumerate(images):
                if i % cols == 0:
                    row = table.add_row().cells
                
                cell = row[i % cols]
                if _try_add_image(cell.paragraphs[0], img_bytes, width_inches=IMAGE_WIDTH_INCHES):
                    pass  # Image added successfully
                else:
                    cell.text = "(afbeelding kon niet worden ingeladen)"
            
            doc.add_paragraph()
        
        # Add spacing between steps
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Error adding step {idx}: {e}")
        raise


def build_workbook_docx_front_and_steps(
    meta: dict[str, Any],
    steps: list[dict[str, Any]]
) -> io.BytesIO:
    """
    Build complete workbook DOCX document.
    
    Args:
        meta: Metadata dictionary with:
            - vak: Subject (BWI/PIE/MVI)
            - opdracht_titel: Assignment title
            - profieldeel: Profile section
            - docent: Teacher name
            - duur: Duration
            - include_materiaalstaat: Include materials list
            - materialen: List of materials
            - cover_bytes: Cover image bytes (optional)
        steps: List of step dictionaries with:
            - title: Step title
            - text_blocks: List of text paragraphs
            - images: List of image bytes
    
    Returns:
        BytesIO object containing DOCX document
        
    Raises:
        Exception: If document generation fails
    """
    try:
        doc = Document()
        
        # Add front matter
        _add_title(doc, meta.get("opdracht_titel") or DEFAULT_TITLE)
        _add_cover(doc, meta)
        _add_meta_block(doc, meta)
        _add_materiaalstaat(doc, meta)
        
        # Add steps
        if steps:
            doc.add_paragraph().add_run(STAPPEN_LABEL).bold = True
            doc.add_paragraph()
            
            for i, step in enumerate(steps, start=1):
                _add_step(doc, i, step)
        
        # Export to BytesIO
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        
        logger.info(f"Workbook generated successfully with {len(steps)} steps")
        return out
    
    except Exception as e:
        logger.error(f"Error building workbook: {e}", exc_info=True)
        raise
