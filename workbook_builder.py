import io
import os
from typing import Optional, Dict
from docx import Document


def _replace_text_everywhere(doc: Document, replacements: Dict[str, str]) -> int:
    """
    Vervang placeholders in:
    - paragraphs
    - tables (alle cellen)
    Retourneert aantal vervangingen (ruwe telling per match).
    """
    replaced_count = 0

    def replace_in_paragraph(paragraph):
        nonlocal replaced_count
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text

        for k, v in replacements.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
                replaced_count += 1

        if new_text != full_text:
            # reset runs en zet alles in 1 run (simpele aanpak)
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = new_text

    # gewone paragrafen
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    return replaced_count


def _template_path_for_vak(vak: str, base_dir: Optional[str] = None) -> str:
    """
    Bepaalt welk templatebestand gebruikt wordt.
    Verwacht templates in: <project>/templates_docx/<VAK>_template.docx
    """
    vak_norm = (vak or "").strip().upper()
    if vak_norm not in {"BWI", "PIE", "MVI"}:
        vak_norm = "BWI"

    if base_dir is None:
        # map waar deze file staat (workbook_builder.py) -> project root
        base_dir = os.path.dirname(os.path.abspath(__file__))

    return os.path.join(base_dir, "templates_docx", f"{vak_norm}_template.docx")


def build_workbook_basic_from_template(meta: dict) -> io.BytesIO:
    """
    Basis werkboekje:
    - Laadt Word template op basis van vak (header/footer/achtergrond komt uit template)
    - Vult placeholders als ze bestaan
    - Als placeholders niet bestaan: zet een klein blokje bovenaan het document
    """
    vak = meta.get("vak", "BWI")
    template_path = _template_path_for_vak(vak)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template niet gevonden: {template_path}")

    doc = Document(template_path)

    titel = (meta.get("opdracht_titel") or "").strip()
    docent = (meta.get("docent") or "").strip()
    duur = (meta.get("duur") or "").strip()
    vak_norm = (vak or "BWI").strip().upper()

    # Placeholders die jij eventueel in je templates kunt zetten (mag ook leeg blijven)
    replacements = {
        "{{TITEL}}": titel,
        "{{DOCENT}}": docent,
        "{{DUUR}}": duur,
        "{{VAK}}": vak_norm,
    }

    replaced = _replace_text_everywhere(doc, replacements)

    # Als er géén placeholders zijn in je template, zetten we minimaal een netjes blok bovenaan
    if replaced == 0:
        # Voeg bovenaan toe: titel + basisinfo (we zetten het aan het begin)
        # python-docx kan niet echt "insert at top" zonder truc; dit is een nette workaround:
        first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        first_para.insert_paragraph_before(titel or "Werkboekje")
        doc.paragraphs[0].style = doc.styles["Title"] if "Title" in [s.name for s in doc.styles] else doc.paragraphs[0].style

        if vak_norm or docent or duur:
            info = []
            if vak_norm:
                info.append(f"Vak: {vak_norm}")
            if docent:
                info.append(f"Docent: {docent}")
            if duur:
                info.append(f"Duur: {duur}")
            first_para.insert_paragraph_before(" · ".join(info))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# Backwards compatible naam (zodat je app.py straks kan blijven uitbreiden)
def build_workbook_docx_front_and_steps(meta: dict, steps: list[dict]) -> io.BytesIO:
    """
    Voor nu: alleen basis vanuit template.
    Later kunnen we hier steps/materiaalstaat/etc. in toevoegen.
    """
    return build_workbook_basic_from_template(meta)

